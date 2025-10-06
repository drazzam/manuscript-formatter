"""
Journal Formatter - Professional Medical Journal Formatting
Based on JAMA, NEJM, BMJ standards for academic medical publications
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from PIL import Image
from typing import Dict, List

class JournalFormatter:
    """
    Applies professional medical journal formatting to manuscripts.
    
    Standards based on:
    - Times New Roman 12pt body text
    - Double spacing
    - 1-inch margins
    - Structured sections
    - Centered figures and tables
    """
    
    # Journal formatting constants
    FONT_NAME = "Times New Roman"
    MARGIN_INCHES = 1.0
    
    def __init__(self, font_size: int = 12, line_spacing: str = "Double", 
                 figure_width: float = 6.0):
        """
        Initialize formatter with custom settings.
        
        Args:
            font_size: Body text font size (10, 11, or 12)
            line_spacing: Line spacing ("Single", "1.5 lines", or "Double")
            figure_width: Maximum figure width in inches
        """
        self.font_size = font_size
        self.line_spacing = self._convert_line_spacing(line_spacing)
        self.figure_width = figure_width
        
    def _convert_line_spacing(self, spacing_str: str) -> float:
        """Convert spacing string to numeric value."""
        spacing_map = {
            "Single": 1.0,
            "1.5 lines": 1.5,
            "Double": 2.0
        }
        return spacing_map.get(spacing_str, 2.0)
    
    def format_document(self, content: Dict, citations: List[Dict], 
                       processor) -> Document:
        """
        Create formatted document with all content properly placed.
        
        Args:
            content: Extracted content dictionary
            citations: List of detected citations
            processor: DocumentProcessor instance for accessing original content
            
        Returns:
            Formatted Document object
        """
        doc = Document()
        
        # Set up document margins
        self._set_margins(doc)
        
        # Add title page
        self._add_title_page(doc, content)
        
        # Add abstract
        self._add_abstract(doc, content)
        
        # Add main body with figures and tables at citation points
        self._add_body_with_citations(doc, content, citations, processor)
        
        # Add references
        self._add_references(doc, content)
        
        return doc
    
    def _set_margins(self, doc: Document):
        """Set 1-inch margins on all sides."""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.MARGIN_INCHES)
            section.bottom_margin = Inches(self.MARGIN_INCHES)
            section.left_margin = Inches(self.MARGIN_INCHES)
            section.right_margin = Inches(self.MARGIN_INCHES)
    
    def _add_title_page(self, doc: Document, content: Dict):
        """Add formatted title page."""
        # Title
        if content['title']:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(content['title'])
            run.font.name = self.FONT_NAME
            run.font.size = Pt(16)
            run.font.bold = True
            title_para.paragraph_format.space_after = Pt(24)
        
        # Authors
        if content['authors']:
            for author in content['authors']:
                author_para = doc.add_paragraph()
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = author_para.add_run(author)
                run.font.name = self.FONT_NAME
                run.font.size = Pt(11)
                author_para.paragraph_format.space_after = Pt(6)
        
        # Page break
        doc.add_page_break()
    
    def _add_abstract(self, doc: Document, content: Dict):
        """Add formatted abstract section."""
        if content['abstract']:
            # Abstract heading
            heading = doc.add_paragraph()
            run = heading.add_run("Abstract")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(14)
            run.font.bold = True
            heading.paragraph_format.space_after = Pt(12)
            
            # Abstract text
            abstract_para = doc.add_paragraph()
            run = abstract_para.add_run(content['abstract'])
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.font_size)
            self._set_paragraph_spacing(abstract_para)
            
            # Space after abstract
            doc.add_paragraph()
    
    def _add_body_with_citations(self, doc: Document, content: Dict, 
                                 citations: List[Dict], processor):
        """
        Add body text with figures and tables inserted at citation points.
        
        NOTE: python-docx limitation - figures are placed as separate paragraphs
        after the paragraph containing the citation, not inline within text.
        This is a known constraint of the library.
        """
        # Create insertion map: paragraph_index -> list of items to insert after
        insertion_map = {}
        for citation in citations:
            para_idx = citation['position']
            if para_idx not in insertion_map:
                insertion_map[para_idx] = []
            insertion_map[para_idx].append(citation)
        
        # Process body paragraphs
        for idx, para_dict in enumerate(content['body']):
            # Add paragraph text
            if para_dict['is_heading']:
                para = doc.add_heading(para_dict['text'], level=1)
                para.runs[0].font.name = self.FONT_NAME
            else:
                para = doc.add_paragraph()
                run = para.add_run(para_dict['text'])
                run.font.name = self.FONT_NAME
                run.font.size = Pt(self.font_size)
                self._set_paragraph_spacing(para)
            
            # Insert figures/tables after this paragraph if citations exist
            if idx in insertion_map:
                for citation in insertion_map[idx]:
                    if citation['type'] == 'figure':
                        self._insert_figure(doc, content, citation, processor)
                    elif citation['type'] == 'table':
                        self._insert_table(doc, content, citation)
    
    def _insert_figure(self, doc: Document, content: Dict, citation: Dict, 
                      processor):
        """
        Insert figure at current position.
        
        Args:
            doc: Document object
            content: Content dictionary
            citation: Citation information
            processor: DocumentProcessor instance
        """
        try:
            # Find matching figure
            fig_num_str = citation['number'].rstrip('ABCDEFGHIJKLMNOPQRSTUVWXYZ')
            fig_num = int(fig_num_str)
            
            matching_figures = [f for f in content['figures'] if f['number'] == fig_num]
            
            if not matching_figures:
                return
            
            figure = matching_figures[0]
            
            # Add spacing before figure
            doc.add_paragraph()
            
            # Add figure
            image_data = figure['image_data']
            image_stream = io.BytesIO(image_data)
            
            # Calculate size while maintaining aspect ratio
            try:
                img = Image.open(io.BytesIO(image_data))
                img_width, img_height = img.size
                aspect_ratio = img_height / img_width
                
                # Set width and calculate height
                width_inches = min(self.figure_width, 6.5)  # Max 6.5 inches
                height_inches = width_inches * aspect_ratio
                
                # Add image
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_stream, width=Inches(width_inches))
                
            except Exception as e:
                # Fallback: add with default size
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_stream, width=Inches(self.figure_width))
            
            # Add caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_text = f"Figure {fig_num}: {figure['caption']}"
            run = caption_para.add_run(caption_text)
            run.font.name = self.FONT_NAME
            run.font.size = Pt(10)
            run.font.bold = True
            
            # Add spacing after figure
            doc.add_paragraph()
            
        except Exception as e:
            # If figure insertion fails, add a placeholder
            para = doc.add_paragraph()
            run = para.add_run(f"[Figure {citation['number']} - insertion failed]")
            run.font.italic = True
    
    def _insert_table(self, doc: Document, content: Dict, citation: Dict):
        """
        Insert table at current position.
        
        Args:
            doc: Document object
            content: Content dictionary
            citation: Citation information
        """
        try:
            table_num = int(citation['number'])
            matching_tables = [t for t in content['tables'] if t['number'] == table_num]
            
            if not matching_tables:
                return
            
            table_data_dict = matching_tables[0]
            table_data = table_data_dict['data']
            
            if not table_data:
                return
            
            # Add spacing before table
            doc.add_paragraph()
            
            # Add caption before table
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_text = f"Table {table_num}: {table_data_dict['caption']}"
            run = caption_para.add_run(caption_text)
            run.font.name = self.FONT_NAME
            run.font.size = Pt(10)
            run.font.bold = True
            
            # Create table
            rows = len(table_data)
            cols = max(len(row) for row in table_data) if table_data else 0
            
            if rows > 0 and cols > 0:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # Populate table
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            cell = table.rows[i].cells[j]
                            cell.text = cell_data
                            
                            # Format cell text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = self.FONT_NAME
                                    run.font.size = Pt(10)
                                
                                # Bold first row (header)
                                if i == 0:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing after table
            doc.add_paragraph()
            
        except Exception as e:
            # If table insertion fails, add a placeholder
            para = doc.add_paragraph()
            run = para.add_run(f"[Table {citation['number']} - insertion failed]")
            run.font.italic = True
    
    def _add_references(self, doc: Document, content: Dict):
        """Add formatted references section."""
        if content['references']:
            # Page break before references
            doc.add_page_break()
            
            # References heading
            heading = doc.add_paragraph()
            run = heading.add_run("References")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(14)
            run.font.bold = True
            heading.paragraph_format.space_after = Pt(12)
            
            # Add each reference
            for ref in content['references']:
                ref_para = doc.add_paragraph()
                run = ref_para.add_run(ref)
                run.font.name = self.FONT_NAME
                run.font.size = Pt(10)
                ref_para.paragraph_format.space_after = Pt(6)
                
                # Hanging indent for references
                ref_para.paragraph_format.left_indent = Inches(0.5)
                ref_para.paragraph_format.first_line_indent = Inches(-0.5)
    
    def _set_paragraph_spacing(self, paragraph):
        """Apply standard spacing to paragraph."""
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = self.line_spacing
        paragraph.paragraph_format.space_after = Pt(0)
