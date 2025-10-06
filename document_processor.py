"""
Document Processor - Core DOCX Processing Engine
Handles extraction of all document elements with robust error handling
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.xmlchemy import OxmlElement
import re
from typing import Dict, List, Tuple, Optional
import io
from PIL import Image

class DocumentProcessor:
    """
    Processes DOCX files to extract and organize content.
    Handles text, figures, tables, and metadata extraction.
    """
    
    def __init__(self, docx_file):
        """
        Initialize processor with DOCX file.
        
        Args:
            docx_file: File path or BytesIO object containing DOCX
        """
        try:
            self.doc = Document(docx_file)
            self.content = {
                'title': '',
                'authors': [],
                'abstract': '',
                'body': [],
                'references': [],
                'figures': [],
                'tables': [],
                'metadata': {}
            }
        except Exception as e:
            raise ValueError(f"Failed to load DOCX file: {str(e)}")
    
    def extract_all_content(self) -> Dict:
        """
        Extract all content from document.
        
        Returns:
            Dictionary containing all extracted content
        """
        try:
            self._extract_title()
            self._extract_authors()
            self._extract_abstract()
            self._extract_body()
            self._extract_references()
            self._extract_figures()
            self._extract_tables()
            return self.content
        except Exception as e:
            raise RuntimeError(f"Content extraction failed: {str(e)}")
    
    def _extract_title(self):
        """Extract document title (usually first bold/large text)."""
        for para in self.doc.paragraphs[:10]:  # Check first 10 paragraphs
            if para.text.strip():
                # Check if text is bold or has "Title:" prefix
                if (para.runs and para.runs[0].bold) or \
                   'title:' in para.text.lower() or \
                   (para.style and 'Title' in para.style.name):
                    self.content['title'] = para.text.replace('Title:', '').strip()
                    return
        
        # Fallback: use first non-empty paragraph
        for para in self.doc.paragraphs:
            if para.text.strip():
                self.content['title'] = para.text.strip()
                return
    
    def _extract_authors(self):
        """Extract author information."""
        author_section = False
        authors = []
        
        for para in self.doc.paragraphs[:30]:  # Check first 30 paragraphs
            text = para.text.strip()
            
            # Detect author section
            if re.search(r'author|affiliation', text, re.IGNORECASE):
                author_section = True
                continue
            
            # Extract authors (look for email patterns, superscripts, etc.)
            if author_section:
                if '@' in text or re.search(r'\d+\s*,', text):
                    authors.append(text)
                elif 'abstract' in text.lower():
                    break
        
        self.content['authors'] = authors
    
    def _extract_abstract(self):
        """Extract abstract section."""
        abstract_text = []
        in_abstract = False
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            
            # Start of abstract
            if re.match(r'^abstract:?$', text, re.IGNORECASE):
                in_abstract = True
                continue
            
            # End of abstract (when we hit Introduction or Methods)
            if in_abstract and re.match(r'^(introduction|methods|keywords):?$', text, re.IGNORECASE):
                break
            
            if in_abstract and text:
                abstract_text.append(text)
        
        self.content['abstract'] = ' '.join(abstract_text)
    
    def _extract_body(self):
        """Extract main body text with section headers."""
        body_content = []
        in_body = False
        skip_patterns = [
            r'^(title|authors?|abstract|keywords?):',
            r'^figure \d+:',
            r'^table \d+:',
            r'^references?:?$'
        ]
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            
            # Start capturing after abstract
            if re.match(r'^(introduction|background|methods)', text, re.IGNORECASE):
                in_body = True
            
            # Stop at references
            if in_body and re.match(r'^references?:?$', text, re.IGNORECASE):
                break
            
            # Skip figure/table legends and other non-body content
            if any(re.match(pattern, text, re.IGNORECASE) for pattern in skip_patterns):
                continue
            
            if in_body and text:
                # Preserve paragraph structure
                body_content.append({
                    'text': text,
                    'style': para.style.name if para.style else 'Normal',
                    'is_heading': self._is_heading(para)
                })
        
        self.content['body'] = body_content
    
    def _is_heading(self, paragraph) -> bool:
        """Determine if paragraph is a heading."""
        if paragraph.style:
            return 'Heading' in paragraph.style.name
        
        # Check if text is bold and short (likely a heading)
        if paragraph.runs and paragraph.runs[0].bold and len(paragraph.text) < 100:
            return True
        
        return False
    
    def _extract_references(self):
        """Extract references section."""
        references = []
        in_references = False
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            
            # Start of references
            if re.match(r'^references?:?$', text, re.IGNORECASE):
                in_references = True
                continue
            
            # Skip if we hit figures/tables section after references
            if in_references and re.match(r'^(figure|table)', text, re.IGNORECASE):
                break
            
            if in_references and text:
                references.append(text)
        
        self.content['references'] = references
    
    def _extract_figures(self):
        """Extract figures (images) from document."""
        figures = []
        figure_counter = 0
        
        for rel_id, rel in self.doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    
                    # Try to find caption near image
                    caption = self._find_figure_caption(figure_counter + 1)
                    
                    figures.append({
                        'number': figure_counter + 1,
                        'image_data': image_data,
                        'caption': caption,
                        'rel_id': rel_id
                    })
                    figure_counter += 1
                except Exception as e:
                    print(f"Warning: Could not extract image {figure_counter + 1}: {str(e)}")
                    continue
        
        self.content['figures'] = figures
    
    def _find_figure_caption(self, figure_num: int) -> str:
        """Find caption for a specific figure number."""
        pattern = rf'^\*?\*?figure\s+{figure_num}[\.:)\s]+(.+)'
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return f"Figure {figure_num}"
    
    def _extract_tables(self):
        """Extract tables from document."""
        tables_data = []
        
        for idx, table in enumerate(self.doc.tables):
            try:
                # Extract table data
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                # Find caption
                caption = self._find_table_caption(idx + 1)
                
                tables_data.append({
                    'number': idx + 1,
                    'data': table_data,
                    'caption': caption,
                    'table_obj': table
                })
            except Exception as e:
                print(f"Warning: Could not extract table {idx + 1}: {str(e)}")
                continue
        
        self.content['tables'] = tables_data
    
    def _find_table_caption(self, table_num: int) -> str:
        """Find caption for a specific table number."""
        pattern = rf'^\*?\*?table\s+{table_num}[\.:)\s]+(.+)'
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return f"Table {table_num}"
    
    def detect_citations(self) -> List[Dict]:
        """
        Detect all figure and table citations in the body text.
        
        Returns:
            List of citation dictionaries with type, number, and context
        """
        citations = []
        
        # Citation patterns
        figure_patterns = [
            r'\(\s*\*?\*?Fig\.?\s+(\d+[A-Z]?)\s*\)',  # (Fig. 1), (Fig. 1A)
            r'\*?\*?Figure\s+(\d+[A-Z]?)\*?\*?',       # Figure 1, **Figure 1A**
            r'\(\s*\*?\*?Figure\s+(\d+[A-Z]?)\s*\)',   # (Figure 1)
        ]
        
        table_patterns = [
            r'\(\s*\*?\*?Tab\.?\s+(\d+)\s*\)',         # (Tab. 1)
            r'\*?\*?Table\s+(\d+)\*?\*?',              # Table 1, **Table 1**
            r'\(\s*\*?\*?Table\s+(\d+)\s*\)',          # (Table 1)
        ]
        
        for para_idx, para_dict in enumerate(self.content['body']):
            text = para_dict['text']
            
            # Find figure citations
            for pattern in figure_patterns:
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    citations.append({
                        'type': 'figure',
                        'number': match.group(1),
                        'position': para_idx,
                        'context': text,
                        'match_text': match.group(0)
                    })
            
            # Find table citations
            for pattern in table_patterns:
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    citations.append({
                        'type': 'table',
                        'number': match.group(1),
                        'position': para_idx,
                        'context': text,
                        'match_text': match.group(0)
                    })
        
        # Remove duplicates and sort by position
        unique_citations = []
        seen = set()
        for citation in citations:
            key = (citation['type'], citation['number'], citation['position'])
            if key not in seen:
                seen.add(key)
                unique_citations.append(citation)
        
        return sorted(unique_citations, key=lambda x: x['position'])
    
    def get_image_size(self, image_data: bytes) -> Tuple[int, int]:
        """
        Get dimensions of an image.
        
        Args:
            image_data: Binary image data
            
        Returns:
            Tuple of (width, height)
        """
        try:
            image = Image.open(io.BytesIO(image_data))
            return image.size
        except Exception:
            return (800, 600)  # Default size if detection fails
